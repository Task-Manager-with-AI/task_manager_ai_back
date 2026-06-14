import base64
import logging
from dataclasses import dataclass, field
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, Optional

import httpx
from docx import Document

from app.core.config import settings
from app.schemas.docx import DocxJobRequest

logger = logging.getLogger(__name__)


@dataclass
class JobState:
    provider_job_id: str
    status: str = "queued"
    error_message: Optional[str] = None
    created_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    started_at: Optional[datetime] = None
    finished_at: Optional[datetime] = None


_jobs: Dict[str, JobState] = {}


def get_job(provider_job_id: str) -> Optional[JobState]:
    return _jobs.get(provider_job_id)


def create_job(payload: DocxJobRequest) -> JobState:
    provider_job_id = f"docx-{payload.jobId}"
    state = JobState(provider_job_id=provider_job_id)
    _jobs[provider_job_id] = state
    return state


async def process_job(payload: DocxJobRequest, state: JobState) -> None:
    try:
        state.status = "processing"
        state.started_at = datetime.now(timezone.utc)
        await _send_callback(payload, state)

        if payload.type == "IMPORT_DOCX":
            result = _import_docx(payload)
        else:
            result = _export_docx(payload)

        state.status = "completed"
        state.finished_at = datetime.now(timezone.utc)
        await _send_callback(payload, state, result=result)
    except Exception as exc:  # noqa: BLE001
        logger.exception("DOCX conversion failed for job=%s", payload.jobId)
        state.status = "failed"
        state.error_message = str(exc)
        state.finished_at = datetime.now(timezone.utc)
        await _send_callback(payload, state)


def _import_docx(payload: DocxJobRequest) -> Dict[str, Any]:
    if not payload.inputFileBase64:
        raise ValueError("inputFileBase64 is required for IMPORT_DOCX")

    binary = base64.b64decode(payload.inputFileBase64)
    doc = Document(BytesIO(binary))

    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs]
    text = "\n".join(line for line in paragraphs if line)

    if len(text) > settings.DOCX_MAX_CONTENT_CHARS:
        text = text[: settings.DOCX_MAX_CONTENT_CHARS]

    return {
        "plainText": text,
        "metadata": {
            "paragraphCount": len(paragraphs),
            "charCount": len(text),
        },
    }


def _export_docx(payload: DocxJobRequest) -> Dict[str, Any]:
    text = payload.sourcePlainText or ""

    if len(text) > settings.DOCX_MAX_CONTENT_CHARS:
        text = text[: settings.DOCX_MAX_CONTENT_CHARS]

    doc = Document()
    for line in text.splitlines() or [""]:
        doc.add_paragraph(line)

    memory_file = BytesIO()
    doc.save(memory_file)
    binary = memory_file.getvalue()

    output_name = payload.requestedFileName or payload.sourceTitle or "document.docx"
    if not output_name.lower().endswith(".docx"):
        output_name = f"{output_name}.docx"

    return {
        "outputFileName": output_name,
        "outputMimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "outputContentBase64": base64.b64encode(binary).decode("utf-8"),
        "metadata": {
            "charCount": len(text),
            "byteSize": len(binary),
        },
    }


async def _send_callback(
    payload: DocxJobRequest,
    state: JobState,
    result: Optional[Dict[str, Any]] = None,
) -> None:
    body: Dict[str, Any] = {
        "documentId": payload.documentId,
        "status": state.status.upper(),
        "providerJobId": state.provider_job_id,
        "errorMessage": state.error_message,
        "startedAt": state.started_at.isoformat().replace("+00:00", "Z") if state.started_at else None,
        "finishedAt": state.finished_at.isoformat().replace("+00:00", "Z") if state.finished_at else None,
    }

    if result is not None:
        body["result"] = result

    async with httpx.AsyncClient(timeout=settings.DOCX_CALLBACK_TIMEOUT_SECONDS) as client:
        response = await client.post(
            payload.callbackUrl,
            json=body,
            headers={"x-docx-callback-secret": payload.callbackSecret},
        )
        response.raise_for_status()
